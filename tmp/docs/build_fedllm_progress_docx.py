from __future__ import annotations

import csv
import math
import os
import re
from collections import OrderedDict
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:/code/Projects/FedLLM")
DOCS = ROOT / "docs"
RUNS = ROOT / "log" / "runs"
OUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs"
IMG_DIR = TMP_DIR / "figures"
OUT_DOCX = OUT_DIR / "FedLLM_LRB_阶段性工作进展与实验结果汇总_表格段落版_20260505.docx"


PY_FONT = r"C:/Windows/Fonts/msyh.ttc"
PY_FONT_BOLD = r"C:/Windows/Fonts/msyhbd.ttc"
if not Path(PY_FONT_BOLD).exists():
    PY_FONT_BOLD = PY_FONT


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fnum(value: str | float | int | None, default: float = math.nan) -> float:
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return float(value)
    s = str(value).strip()
    if not s or s.lower() in {"n/a", "na", "nan", "pending"}:
        return default
    try:
        return float(s)
    except ValueError:
        return default


def fmt_float(value: str | float | int | None, digits: int = 4) -> str:
    v = fnum(value)
    if math.isnan(v):
        return "n/a"
    return f"{v:.{digits}f}"


def fmt_pct(value: str | float | int | None, digits: int = 2) -> str:
    v = fnum(value)
    if math.isnan(v):
        return "n/a"
    return f"{v * 100:.{digits}f}%"


def seconds_to_hms(value: str | float | int | None) -> str:
    v = fnum(value)
    if math.isnan(v):
        return "n/a"
    seconds = int(round(v))
    h, rem = divmod(seconds, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def method_label(defense: str, param: str | float | int | None = None) -> str:
    if defense == "none":
        return "none"
    if param is None or str(param).strip().lower() in {"", "n/a", "nan"}:
        return defense
    p = str(param).strip()
    try:
        x = float(p)
        if defense == "compression":
            p = str(int(x))
        elif abs(x) < 0.001 and x != 0:
            p = f"{x:.0e}".replace("e-0", "e-").replace("e+0", "e+")
        else:
            p = f"{x:g}"
    except ValueError:
        pass
    return f"{defense}@{p}"


def load_defense_results() -> list[dict[str, str]]:
    path = (
        RUNS
        / "defense_baselines_sst2_b2_gpt2_20260501_010024"
        / "defense_baselines_sst2_b2_gpt2_20260501_010024"
        / "results.csv"
    )
    return read_csv(path)


def load_ablation_combined() -> list[dict[str, str]]:
    path = RUNS / "lrb_ablation_sst2_b2_gpt2_k0.5_20260501_004737" / "ablation_combined_summary.csv"
    return read_csv(path)


def utility_rows_from_known_runs() -> list[dict[str, str | float]]:
    # Focused utility runs generated at different times all use the same clean anchor.
    # The "none" row is represented once using the common summary values.
    rows: list[dict[str, str | float]] = [
        {
            "method": "none",
            "defense": "none",
            "param": "n/a",
            "eval_accuracy": 0.913226,
            "eval_macro_f1": 0.913184,
            "eval_loss": 0.246637,
            "utility_drop": 0.0,
            "time": "00:42:51",
            "source": "UTILITY_RESULTS_ANALYSIS_20260426.md",
        },
        {
            "method": "topk@0.1",
            "defense": "topk",
            "param": "0.1",
            "eval_accuracy": 0.912462,
            "eval_macro_f1": 0.912430,
            "eval_loss": 0.243324,
            "utility_drop": 0.000764,
            "time": "06:03:08",
            "source": "sst2_b2_gpt2_topk_0.1",
        },
        {
            "method": "compression@8",
            "defense": "compression",
            "param": "8",
            "eval_accuracy": 0.911315,
            "eval_macro_f1": 0.911290,
            "eval_loss": 0.263210,
            "utility_drop": 0.001911,
            "time": "07:07:37",
            "source": "sst2_b2_gpt2_compression_8",
        },
        {
            "method": "mixup@0.3",
            "defense": "mixup",
            "param": "0.3",
            "eval_accuracy": 0.910933,
            "eval_macro_f1": 0.910906,
            "eval_loss": 0.239469,
            "utility_drop": 0.002293,
            "time": "00:48:40",
            "source": "sst2_b2_gpt2_mixup_0.3",
        },
        {
            "method": "lrb@0.2",
            "defense": "lrb",
            "param": "0.2",
            "eval_accuracy": 0.821865,
            "eval_macro_f1": 0.821360,
            "eval_loss": 0.441765,
            "utility_drop": 0.091361,
            "time": "08:52:41",
            "source": "sst2_b2_gpt2_lrb_0.2",
        },
        {
            "method": "noise@5e-4",
            "defense": "noise",
            "param": "5e-4",
            "eval_accuracy": 0.715979,
            "eval_macro_f1": 0.715617,
            "eval_loss": 0.552434,
            "utility_drop": 0.197247,
            "time": "01:54:29",
            "source": "sst2_b2_gpt2_noise_5e-4",
        },
        {
            "method": "dpsgd@5e-4",
            "defense": "dpsgd",
            "param": "5e-4",
            "eval_accuracy": 0.504205,
            "eval_macro_f1": 0.366347,
            "eval_loss": 2.612275,
            "utility_drop": 0.409021,
            "time": "03:29:31",
            "source": "sst2_b2_gpt2_dpsgd_5e-4",
        },
        {
            "method": "lrb@0.35",
            "defense": "lrb",
            "param": "0.35",
            "eval_accuracy": 0.868119,
            "eval_macro_f1": 0.868010,
            "eval_loss": 0.356615,
            "utility_drop": 0.045107,
            "time": "05:29:49",
            "source": "sst2_b2_gpt2_lrb_0.35",
        },
        {
            "method": "lrb@0.5",
            "defense": "lrb",
            "param": "0.5",
            "eval_accuracy": 0.892584,
            "eval_macro_f1": 0.892472,
            "eval_loss": 0.321702,
            "utility_drop": 0.020642,
            "time": "01:56:47",
            "source": "sst_b2_gpt2_lrb_0.5",
        },
        {
            "method": "topk@0.3",
            "defense": "topk",
            "param": "0.3",
            "eval_accuracy": 0.910933,
            "eval_macro_f1": 0.910913,
            "eval_loss": 0.256373,
            "utility_drop": 0.002293,
            "time": "01:16:24",
            "source": "sst2_b2_gpt2_topk_0.3",
        },
        {
            "method": "compression@16",
            "defense": "compression",
            "param": "16",
            "eval_accuracy": 0.909021,
            "eval_macro_f1": 0.909012,
            "eval_loss": 0.245559,
            "utility_drop": 0.004205,
            "time": "02:03:04",
            "source": "sst2_b2_gpt2_compression_16",
        },
    ]
    return rows


def find_row(rows: list[dict[str, str]], defense: str, param: str | None = None) -> dict[str, str] | None:
    for row in rows:
        if row.get("defense") != defense:
            continue
        if param is None:
            return row
        val = row.get("defense_param_value", "")
        if val == param:
            return row
        try:
            if abs(float(val) - float(param)) < 1e-9:
                return row
        except Exception:
            pass
    return None


def summarize_defense_by_family(rows: list[dict[str, str]]) -> list[list[str]]:
    groups: OrderedDict[str, list[dict[str, str]]] = OrderedDict()
    for row in rows:
        groups.setdefault(row["defense"], []).append(row)

    out: list[list[str]] = []
    for defense, vals in groups.items():
        vals_sorted = sorted(vals, key=lambda r: fnum(r.get("defense_param_value"), -1))
        if defense == "none":
            row = vals_sorted[0]
            out.append(
                [
                    "none",
                    "n/a",
                    fmt_float(row["rec_token_mean"], 6),
                    fmt_float(row["agg_r1fm_r2fm"], 3),
                    "clean 更新泄露严重，作为风险锚点",
                ]
            )
        elif defense == "lrb":
            out.append(["lrb", "0.05/0.1/0.2/0.35/0.5", "0.000000", "0.000", "全部 DAGER=0，privacy 已饱和"])
        elif defense == "topk":
            safe = [r for r in vals_sorted if fnum(r["rec_token_mean"]) == 0]
            leak = [r for r in vals_sorted if fnum(r["rec_token_mean"]) > 0]
            safe_params = "/".join(f"{fnum(r['defense_param_value']):g}" for r in safe)
            leak_desc = "; ".join(
                f"{fnum(r['defense_param_value']):g}->{fmt_float(r['rec_token_mean'], 3)}" for r in leak
            )
            out.append(["topk", safe_params, "0.000000", "0.000", f"安全区到 0.3；更大比例开始泄露：{leak_desc}"])
        elif defense == "compression":
            out.append(["compression", "4/8/16", "0.000000", "0.000", "4/8/16 稳定；2bit 在 56/100 SVD 失败；32bit 退化"])
        elif defense == "noise":
            first, last = vals_sorted[0], vals_sorted[-1]
            out.append(
                [
                    "noise",
                    "1e-6 -> 1e-3",
                    f"{fmt_float(first['rec_token_mean'], 3)} -> {fmt_float(last['rec_token_mean'], 3)}",
                    f"{fmt_float(first['agg_r1fm_r2fm'], 3)} -> {fmt_float(last['agg_r1fm_r2fm'], 3)}",
                    "需要较大扰动才有效，token 仍有残留",
                ]
            )
        elif defense == "dpsgd":
            first, last = vals_sorted[0], vals_sorted[-1]
            out.append(
                [
                    "dpsgd",
                    "1e-6 -> 1e-3",
                    f"{fmt_float(first['rec_token_mean'], 3)} -> {fmt_float(last['rec_token_mean'], 3)}",
                    f"{fmt_float(first['agg_r1fm_r2fm'], 3)} -> {fmt_float(last['agg_r1fm_r2fm'], 3)}",
                    "理论基线重要，但当前 utility 代价大",
                ]
            )
        elif defense == "mixup":
            vals2 = [fnum(r["rec_token_mean"]) for r in vals_sorted]
            rouge = [fnum(r["agg_r1fm_r2fm"]) for r in vals_sorted]
            out.append(
                [
                    "mixup",
                    "0.1~2.0",
                    f"{min(vals2):.3f}~{max(vals2):.3f}",
                    f"{min(rouge):.3f}~{max(rouge):.3f}",
                    "utility 友好但 privacy 恶化，不适合作为隐私防御",
                ]
            )
        elif defense == "soteria":
            vals2 = [fnum(r["rec_token_mean"]) for r in vals_sorted]
            rouge = [fnum(r["agg_r1fm_r2fm"]) for r in vals_sorted]
            out.append(
                [
                    "soteria",
                    "10~90",
                    f"{min(vals2):.3f}~{max(vals2):.3f}",
                    f"{min(rouge):.3f}~{max(rouge):.3f}",
                    "迁移到当前 LLM 设置后明显恶化",
                ]
            )
    order = {"none": 0, "lrb": 1, "topk": 2, "compression": 3, "noise": 4, "dpsgd": 5, "mixup": 6, "soteria": 7}
    return sorted(out, key=lambda x: order.get(x[0], 99))


def safe_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", name)


def draw_bar_chart(
    path: Path,
    title: str,
    data: list[tuple[str, float]],
    ylabel: str,
    *,
    max_value: float | None = None,
    lower_is_better: bool = False,
    highlight: set[str] | None = None,
    width: int = 1600,
    height: int = 900,
):
    highlight = highlight or set()
    font_title = ImageFont.truetype(PY_FONT_BOLD, 44)
    font = ImageFont.truetype(PY_FONT, 25)
    font_small = ImageFont.truetype(PY_FONT, 21)
    font_bold = ImageFont.truetype(PY_FONT_BOLD, 25)
    img = Image.new("RGB", (width, height), "#fbfaf7")
    d = ImageDraw.Draw(img)
    left, right, top, bottom = 150, 70, 125, 170
    chart_w = width - left - right
    chart_h = height - top - bottom
    d.text((left, 42), title, fill="#1f2933", font=font_title)
    d.text((left, 88), ylabel, fill="#53606c", font=font_small)
    vals = [v for _, v in data]
    if max_value is None:
        max_value = max(vals) * 1.12 if vals else 1.0
    if max_value == 0:
        max_value = 1.0
    # Grid
    for i in range(6):
        y = top + chart_h - chart_h * i / 5
        d.line((left, y, width - right, y), fill="#e2e8ef", width=2)
        label = f"{max_value * i / 5:.2f}"
        tw = d.textlength(label, font=font_small)
        d.text((left - tw - 16, y - 12), label, fill="#69747f", font=font_small)
    d.line((left, top, left, top + chart_h), fill="#bdc7d1", width=2)
    d.line((left, top + chart_h, width - right, top + chart_h), fill="#bdc7d1", width=2)
    n = len(data)
    gap = 26
    bar_w = max(28, (chart_w - gap * (n + 1)) / n)
    colors = ["#2a9d8f", "#4e79a7", "#f28e2b", "#e15759", "#7f6aa9", "#59a14f", "#edc948", "#76b7b2"]
    for idx, (label, val) in enumerate(data):
        x0 = left + gap + idx * (bar_w + gap)
        x1 = x0 + bar_w
        y1 = top + chart_h
        y0 = y1 - (val / max_value) * chart_h
        color = "#d94841" if label in highlight else colors[idx % len(colors)]
        if lower_is_better and val == min(vals):
            color = "#2a9d8f"
        d.rounded_rectangle((x0, y0, x1, y1), radius=8, fill=color)
        value_label = f"{val:.3f}" if val < 10 else f"{val:.1f}"
        tw = d.textlength(value_label, font=font_small)
        d.text((x0 + bar_w / 2 - tw / 2, y0 - 32), value_label, fill="#26323c", font=font_small)
        # Multi-line tick labels
        parts = label.replace("@", "\n@").split("\n")
        yy = y1 + 20
        for part in parts[:2]:
            f = font_bold if label in highlight else font_small
            tw = d.textlength(part, font=f)
            d.text((x0 + bar_w / 2 - tw / 2, yy), part, fill="#26323c", font=f)
            yy += 28
    img.save(path, quality=95)


def draw_scatter_chart(
    path: Path,
    title: str,
    points: list[dict[str, float | str]],
    *,
    width: int = 1600,
    height: int = 950,
):
    font_title = ImageFont.truetype(PY_FONT_BOLD, 42)
    font = ImageFont.truetype(PY_FONT, 24)
    font_small = ImageFont.truetype(PY_FONT, 20)
    font_bold = ImageFont.truetype(PY_FONT_BOLD, 22)
    img = Image.new("RGB", (width, height), "#fbfaf7")
    d = ImageDraw.Draw(img)
    left, right, top, bottom = 160, 230, 135, 150
    chart_w = width - left - right
    chart_h = height - top - bottom
    d.text((left, 42), title, fill="#1f2933", font=font_title)
    d.text((left, 92), "横轴为 utility drop，纵轴为 DAGER rec_token_mean；左下越好。", fill="#53606c", font=font)
    x_max = max(float(p["drop"]) for p in points) * 1.08
    y_max = max(float(p["rec"]) for p in points) * 1.12
    x_max = max(x_max, 0.02)
    y_max = max(y_max, 0.9)
    # Grid
    for i in range(6):
        x = left + chart_w * i / 5
        d.line((x, top, x, top + chart_h), fill="#e2e8ef", width=2)
        label = f"{x_max * i / 5:.2f}"
        tw = d.textlength(label, font=font_small)
        d.text((x - tw / 2, top + chart_h + 18), label, fill="#69747f", font=font_small)
    for i in range(6):
        y = top + chart_h - chart_h * i / 5
        d.line((left, y, left + chart_w, y), fill="#e2e8ef", width=2)
        label = f"{y_max * i / 5:.2f}"
        tw = d.textlength(label, font=font_small)
        d.text((left - tw - 18, y - 12), label, fill="#69747f", font=font_small)
    d.line((left, top, left, top + chart_h), fill="#bdc7d1", width=2)
    d.line((left, top + chart_h, left + chart_w, top + chart_h), fill="#bdc7d1", width=2)
    d.text((left + chart_w / 2 - 80, height - 58), "Utility drop", fill="#53606c", font=font)
    d.text((35, top + chart_h / 2 - 20), "rec_token", fill="#53606c", font=font)
    colors = {
        "none": "#6b7280",
        "lrb": "#d94841",
        "topk": "#2a9d8f",
        "compression": "#4e79a7",
        "noise": "#f28e2b",
        "dpsgd": "#7f6aa9",
        "mixup": "#edc948",
        "proj": "#8b5cf6",
    }
    zero_cluster = {
        "proj_only@0.5",
        "topk@0.1",
        "compression@8",
        "topk@0.3",
        "compression@16",
        "lrb@0.5",
        "lrb@0.35",
    }
    cluster_rows: list[tuple[str, str]] = []
    for p in points:
        x = left + (float(p["drop"]) / x_max) * chart_w
        y = top + chart_h - (float(p["rec"]) / y_max) * chart_h
        family = str(p.get("family", "none"))
        color = colors.get(family, "#333333")
        r = 12 if family in {"lrb", "topk", "compression", "proj"} else 10
        d.ellipse((x - r, y - r, x + r, y + r), fill=color, outline="#ffffff", width=3)
        label = str(p["label"])
        if label in zero_cluster:
            cluster_rows.append((label, family))
            continue
        dx, dy = (14, -30 if y > top + 60 else 14)
        d.text((x + dx, y + dy), label, fill="#26323c", font=font_bold if family in {"lrb", "proj"} else font_small)
    # Dedicated callout for the crowded zero-recovery cluster.
    bx, by, bw, bh = left + 300, top + chart_h - 205, 470, 185
    d.rounded_rectangle((bx, by, bx + bw, by + bh), radius=12, fill="#ffffff", outline="#d8e2ec", width=2)
    d.text((bx + 18, by + 12), "DAGER=0 簇（按 utility drop 递增）", fill="#26323c", font=font_bold)
    zero_order = ["proj_only@0.5", "topk@0.1", "compression@8", "topk@0.3", "compression@16", "lrb@0.5", "lrb@0.35"]
    fam_by_label = {label: fam for label, fam in cluster_rows}
    cx, cy = bx + 20, by + 50
    for idx, label in enumerate(zero_order):
        fam = fam_by_label.get(label)
        if not fam:
            continue
        col = colors.get(fam, "#333333")
        row_x = cx if idx < 4 else bx + 245
        row_y = cy + (idx % 4) * 30
        d.ellipse((row_x, row_y + 5, row_x + 15, row_y + 20), fill=col)
        d.text((row_x + 24, row_y), label, fill="#26323c", font=font_small)
    # Legend
    lx, ly = width - right + 35, top + 10
    for name, label in [("lrb", "LRB"), ("topk", "Top-k"), ("compression", "Compression"), ("proj", "Projection-LRB"), ("none", "None/Other")]:
        d.ellipse((lx, ly, lx + 20, ly + 20), fill=colors[name])
        d.text((lx + 32, ly - 4), label, fill="#26323c", font=font_small)
        ly += 42
    img.save(path, quality=95)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold=False, size=9, color="1f2933", align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, header_fill="D9EAF7", header_color="123047", font_size=8.5, first_col_bold=False):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string("1f2933")
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(header_color)
            elif c_idx == 0 and first_col_bold:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True


def set_table_widths(table, widths_cm: list[float]):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_table_text_version(doc: Document, headers: list[str], rows: list[list[str]]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("表格文字版：")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(9.2)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("123047")

    for idx, row in enumerate(rows, start=1):
        parts = []
        for col_idx, header in enumerate(headers):
            value = row[col_idx] if col_idx < len(row) else ""
            value = str(value).replace("\n", " ").strip()
            parts.append(f"{header}={value}")
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        rr = p.add_run(f"第{idx}行：" + "；".join(parts) + "。")
        rr.font.name = "微软雅黑"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        rr.font.size = Pt(9.0)
        rr.font.color.rgb = RGBColor.from_string("34404a")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None, font_size=8.5, text_version: bool = True):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) < 18 and i != len(row) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(val), size=font_size, align=align)
    style_table(table, font_size=font_size, first_col_bold=True)
    if widths_cm:
        set_table_widths(table, widths_cm)
    if text_version:
        add_table_text_version(doc, headers, [[str(v) for v in row] for row in rows])
    doc.add_paragraph()
    return table


def set_doc_defaults(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        st = styles[style_name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.color.rgb = RGBColor.from_string("123047")
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(13.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.85)
        section.right_margin = Cm(1.85)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc: Document, text: str, *, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(item)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10.2)


def add_note_box(doc: Document, title: str, lines: list[str], fill="F3F8FB"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("123047")
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        rr = p.add_run(line)
        rr.font.name = "微软雅黑"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        rr.font.size = Pt(9.8)
        rr.font.color.rgb = RGBColor.from_string("1f2933")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_text_version(doc, ["标题", "内容"], [[title, line] for line in lines])
    doc.add_paragraph()


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("53606c")


def add_image(doc: Document, path: Path, width_cm: float, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def build_charts(defense_rows: list[dict[str, str]], utility_rows: list[dict[str, str | float]], ablation_rows: list[dict[str, str]]) -> dict[str, Path]:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    # Privacy bars for representative DAGER results.
    reps = [
        ("none", find_row(defense_rows, "none")),
        ("topk@0.1", find_row(defense_rows, "topk", "0.1")),
        ("compression@8", find_row(defense_rows, "compression", "8")),
        ("lrb@0.5", find_row(defense_rows, "lrb", "0.5")),
        ("noise@5e-4", find_row(defense_rows, "noise", "0.0005")),
        ("dpsgd@5e-4", find_row(defense_rows, "dpsgd", "0.0005")),
        ("mixup@0.3", find_row(defense_rows, "mixup", "0.3")),
        ("soteria@30", find_row(defense_rows, "soteria", "30")),
    ]
    privacy_data = [(label, fnum(row["rec_token_mean"]) if row else 0.0) for label, row in reps]
    privacy_chart = IMG_DIR / "privacy_rec_token_bar.png"
    draw_bar_chart(
        privacy_chart,
        "DAGER 文本恢复风险：代表性 defense 的 rec_token_mean",
        privacy_data,
        "越低越好；0 表示当前 DAGER 未恢复出 token",
        max_value=1.05,
        lower_is_better=True,
        highlight={"lrb@0.5"},
    )
    # Utility drops.
    util_order = ["topk@0.1", "compression@8", "topk@0.3", "compression@16", "lrb@0.5", "lrb@0.35", "lrb@0.2", "noise@5e-4", "dpsgd@5e-4"]
    util_map = {str(r["method"]): r for r in utility_rows}
    util_data = [(m, float(util_map[m]["utility_drop"])) for m in util_order if m in util_map]
    utility_chart = IMG_DIR / "utility_drop_bar.png"
    draw_bar_chart(
        utility_chart,
        "同等/近似 DAGER privacy 下的 utility drop",
        util_data,
        "相对 clean accuracy 的下降；越低越好",
        max_value=0.43,
        lower_is_better=True,
        highlight={"lrb@0.5"},
    )
    # Privacy-utility scatter, including ablation projection point.
    rec_map = {}
    for row in defense_rows:
        label = method_label(row["defense"], row.get("defense_param_value"))
        rec_map[label] = fnum(row["rec_token_mean"])
    points: list[dict[str, float | str]] = []
    for row in utility_rows:
        method = str(row["method"])
        if method == "none":
            rec = rec_map.get("none", 0.833506)
        else:
            rec = rec_map.get(method, 0.0 if method in {"topk@0.1", "topk@0.3", "compression@8", "compression@16", "lrb@0.2", "lrb@0.35", "lrb@0.5"} else math.nan)
        if math.isnan(rec):
            continue
        family = str(row["defense"])
        points.append({"label": method, "drop": float(row["utility_drop"]), "rec": rec, "family": family})
    # Add projection-LRB from ablation as current method candidate.
    proj = next(r for r in ablation_rows if r["variant"] == "proj_only")
    points.append(
        {
            "label": "proj_only@0.5",
            "drop": max(0.0, fnum(proj["utility_drop"])),
            "rec": fnum(proj["rec_token_mean"]),
            "family": "proj",
        }
    )
    scatter_chart = IMG_DIR / "privacy_utility_scatter.png"
    draw_scatter_chart(scatter_chart, "Privacy-Utility tradeoff：当前事实点与新候选", points)
    # Ablation accuracy bar.
    order = ["none", "identity_lrb", "clip_only", "proj_only", "proj_clip", "full_lrb", "pool_full", "rule_only", "empirical_only", "uniform_all_sensitive"]
    ab_map = {r["variant"]: r for r in ablation_rows}
    ab_labels = {
        "identity_lrb": "identity",
        "clip_only": "clip",
        "proj_only": "proj_only",
        "proj_clip": "proj+clip",
        "full_lrb": "full_lrb",
        "pool_full": "pool",
        "rule_only": "rule",
        "empirical_only": "empirical",
        "uniform_all_sensitive": "uniform",
    }
    ab_data = [(ab_labels.get(v, v), fnum(ab_map[v]["eval_accuracy"])) for v in order if v in ab_map]
    ablation_chart = IMG_DIR / "ablation_accuracy_bar.png"
    draw_bar_chart(
        ablation_chart,
        "LRB 消融：在 DAGER=0 的变体中谁保留 utility",
        ab_data,
        "eval_accuracy；越高越好",
        max_value=0.94,
        lower_is_better=False,
        highlight={"proj_only", "full_lrb"},
    )
    return {
        "privacy": privacy_chart,
        "utility": utility_chart,
        "scatter": scatter_chart,
        "ablation": ablation_chart,
    }


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("FedLLM / LRB 阶段性工作进展与实验结果汇总")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(25)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("123047")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("基于现有分析文档、DAGER defense baselines、utility runs 与 LRB ablation runs")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor.from_string("53606c")

    meta = [
        ("整理日期", "2026-05-05"),
        ("主数据集 / 模型", "SST2 / GPT-2 fine-tuned checkpoint"),
        ("主攻击面", "Full-gradient DAGER"),
        ("主方法候选", "Projection-LRB / LRB-lite（由消融结果支持）"),
        ("主要日志来源", "defense_baselines_sst2_b2_gpt2_20260501_010024；utility260426；lrb_ablation_sst2_b2_gpt2_k0.5_20260501_004737"),
    ]
    add_table(doc, ["项目", "内容"], meta, widths_cm=[4, 12.5], font_size=9.5)

    add_note_box(
        doc,
        "核心结论先行",
        [
            "当前工作已经形成 DAGER 攻击、防御 baseline、utility、LRB 机制消融的闭环证据，但强证据边界主要是 SST2+GPT2+batch=2 的 full-gradient DAGER。",
            "Clean FedSGD 泄露严重；LRB、topk、compression 均可在若干点上把当前 DAGER 恢复打到 0，但这不等价于形式化隐私保证。",
            "仅看 full-gradient DAGER，topk@0.1 与 compression@8 仍是更强经验 tradeoff；完整 full_lrb@0.5 已改善 utility，但仍偏重。",
            "最新消融显示 projection bottleneck 是主效应：proj_only@0.5 在单次消融运行中达到 DAGER=0 并保持 clean-level utility，因此下一阶段应把主方法收束为 Projection-LRB。",
        ],
        fill="EAF5F2",
    )
    doc.add_page_break()


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    defense_rows = load_defense_results()
    ablation_rows = load_ablation_combined()
    utility_rows = utility_rows_from_known_runs()
    charts = build_charts(defense_rows, utility_rows, ablation_rows)

    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc)

    add_heading(doc, "一、文档定位与当前工作状态", 1)
    add_body(
        doc,
        "本文参考 docs/初步思路模版.docx 的研究计划结构，将当前 FedLLM / FL-LLM 方向整理为一份阶段性研究进展文档。它不是重新粘贴已有 Markdown，而是把任务价值、文献与 baseline 位置、LRB 方法机制、已完成实验事实、当前判断和后续计划串成一条可继续写论文的主线。",
    )
    add_body(
        doc,
        "截至 2026-05-05，项目已经从“复现 DAGER + 枚举 baseline”推进到更清晰的阶段：full-gradient DAGER 风险已被量化，多个 defense 的 privacy/utility 已经有结果，LRB 的主要组件也通过完整消融开始被拆开。当前最重要的变化是，主线不再是继续证明“DAGER 可以被打掉”，而是要证明“哪一种 recoverability bottleneck 能在相近 privacy 下保留更多效用，并且能泛化到 PEFT / partial-gradient 等更广攻击面”。",
    )
    add_note_box(
        doc,
        "审稿口径下的证据边界",
        [
            "已经完成并可较强支撑的证据：SST2 + GPT2 fine-tuned checkpoint + batch=2 + n_inputs=100 的 full-gradient DAGER attack-time privacy；full training utility fixed points；一次完整 LRB ablation。",
            "尚不能直接外推的证据：LoRA/PEFT 训练期 defense、partial-gradient attack、跨数据集/backbone、多 seed 统计显著性、adaptive attack 与严格 DP accounting。",
            "因此当前论文主张应写成：Projection-LRB 是由 full-gradient DAGER 与消融结果支持的主方法候选，而不是已经在所有 FL-LLM 场景全面优于压缩类 baseline。",
        ],
        fill="FFF6E6",
    )
    add_note_box(
        doc,
        "当前状态一句话",
        [
            "DAGER clean 泄露严重，问题成立；topk/compression 是 full-gradient DAGER 下必须认真比较的强经验 baseline；full_lrb 能强力抑制恢复但过防御；消融显示低分辨率 projection bottleneck 才是当前最值得收束为主方法的机制。",
        ],
    )

    add_heading(doc, "二、背景与动机", 1)
    add_body(
        doc,
        "大模型微调正在从集中式数据汇聚转向联邦学习、跨机构协作训练和 PEFT adapter 共享。这样的设置看起来天然适合隐私场景：医疗文本、金融评论、企业客服、用户输入日志等原始文本不离开本地，只上传梯度、模型更新或轻量参数更新。但梯度反演与文本恢复攻击表明，“不上传原文”并不等价于“原文不可恢复”。在语言模型中，token embedding、layer-wise gradient 和序列损失共同形成了足以暴露样本内容的高分辨率信号。",
    )
    add_body(
        doc,
        "本项目当前的 clean DAGER 结果给出了直接动机：在 SST2 + GPT2 + batch=2 + n_inputs=100 的设置下，未防御 FedSGD 梯度的 rec_token_mean 达到 0.833506，ROUGE-1+ROUGE-2 达到 141.710856。这说明攻击者并不是只能恢复模糊主题，而是能够恢复大量 token 级文本信息。因此，本工作不是为了在一个 toy setting 中展示攻击存在，而是要面向真实 FL-LLM / PEFT 协作训练中的中间更新泄露问题，寻找可落地的防御机制。",
    )
    add_body(
        doc,
        "进一步的动机来自 privacy-utility 冲突：简单把 DAGER 打到 0 并不困难，困难在于保留训练效用。当前实验已经显示，完整 LRB 能稳定抑制 DAGER，但 utility 代价偏高；topk/compression 在 full-gradient DAGER 下很强，却更像通信压缩副作用而不是针对文本 recoverability 的机制解释。论文真正需要回答的是：能否设计一种面向 layer-wise recoverability 的 bottleneck，使更新仍可用于学习任务信号，但不再保留可被 DAGER 等攻击恢复的高分辨率文本细节。",
    )

    add_heading(doc, "三、前人工作局限性", 1)
    add_body(
        doc,
        "已有工作为本项目提供了攻击、隐私和压缩三类重要参照，但各自都留下了适合本工作切入的缺口。当前实验也印证了这些局限：强扰动方法通常牺牲 utility，通信压缩方法在当前 DAGER 下表现强但机制目标不完全对齐，表示扰动和训练侧混合方法在 GPT2/FedSGD 设置中并不稳定。",
    )
    add_table(
        doc,
        ["方向", "代表思路", "主要局限性", "对本工作的启发"],
        [
            [
                "梯度反演攻击",
                "DAGER / LAMP 等从梯度恢复文本",
                "多数工作强调攻击能力，较少给出适合 FL-LLM 训练流程的低损耗防御；对 PEFT、partial-gradient 和 defense 后 utility 的系统闭环不足。",
                "需要把攻击复现、defense、end-to-end utility 和消融放在同一评价链路里。",
            ],
            [
                "噪声与 DP-SGD",
                "对梯度裁剪后加噪，获得形式化 DP 语义",
                "理论上强，但在大模型梯度和小 batch 文本任务中 utility 代价明显；本项目 noise/DP-SGD 结果也显示较难兼顾 privacy 与 accuracy。",
                "不能只追求“加噪后看不见”，而要保留任务梯度中的可学习结构。",
            ],
            [
                "通信压缩",
                "Top-k sparsification、SVD/quantization compression",
                "当前 full-gradient DAGER 下是强 baseline，但核心目标是少传或压缩，不是解释哪些 layer/direction 泄露文本；跨 PEFT/partial-gradient 攻击面仍未验证。",
                "主方法必须正面比较 topk/compression，同时说明 recoverability bottleneck 的机制差异。",
            ],
            [
                "表示扰动与训练增强",
                "Soteria、mixup 等表示层或训练侧防御",
                "迁移到 GPT2/FedSGD 当前设置后不稳定：soteria 几乎完全泄露，mixup 多点 privacy 失败。",
                "需要面向 Transformer 梯度结构重新设计，而不是直接搬用 CV/表示防御经验。",
            ],
            [
                "PEFT/局部梯度泄露",
                "LoRA/adapter 更新反演、partial transformer gradient 攻击",
                "更贴近真实部署，但当前仓库只有 eval-first 框架和威胁建模，尚缺完整 defense 结果。",
                "这是下一阶段证明泛化意义的关键实验面。",
            ],
        ],
        widths_cm=[3.0, 4.0, 5.6, 5.2],
        font_size=7.8,
    )

    add_heading(doc, "四、希望解决的问题", 1)
    add_body(
        doc,
        "本工作希望解决的不是单一实验点上的“攻击是否成功”问题，而是 FL-LLM 中间更新如何在可训练与不可恢复之间取得平衡的问题。具体来说，需要把研究问题拆成四个层次：先证明 clean 更新确实泄露，再建立强 baseline 对照，然后设计可解释的 recoverability bottleneck，最后验证该 bottleneck 是否能跨训练形式和攻击面泛化。",
    )
    add_table(
        doc,
        ["问题层次", "要回答的问题", "当前已有依据", "下一步需要补强"],
        [
            [
                "风险是否成立",
                "客户端只上传梯度/更新时，攻击者是否仍能恢复文本？",
                "clean DAGER rec_token_mean=0.833506，ROUGE-1+ROUGE-2=141.710856。",
                "跨数据集/backbone 复核，避免单一 SST2/GPT2 偶然性。",
            ],
            [
                "防御是否有效",
                "哪些方法能把 token 恢复和 ROUGE 恢复压到接近 0？",
                "LRB@0.05~0.5、topk@0.01~0.3、compression@4/8/16 在 DAGER 下表现强。",
                "补齐稳定点说明；DAGER=0 只能表示当前攻击未恢复，不等价于严格隐私保证。",
            ],
            [
                "效用是否可接受",
                "同等 privacy 下，哪个 defense 保留最多 task accuracy？",
                "topk@0.1 和 compression@8 是当前 full-gradient DAGER 的强经验 tradeoff；full_lrb utility 偏弱。",
                "把 Projection-LRB 作为主方法纳入 main result，与强 baseline 做同等设置和多 seed 比较。",
            ],
            [
                "机制是否清楚",
                "防御效果来自 clipping、projection、noise 还是 layer sensitivity？",
                "ablation 显示 proj_only@0.5 在单次运行中 DAGER=0 且 accuracy=0.915520，是当前最强机制线索。",
                "做 projection-only keep-ratio sweep、多 seed 与细消融，确认最宽松 bottleneck 和方差。",
            ],
            [
                "场景是否泛化",
                "该机制是否适用于 PEFT/LoRA、partial-gradient 和更真实的联邦微调？",
                "PEFT eval-first 框架已有，partial-gradient 威胁已建模。",
                "补 LoRA 训练期 defense、partial-gradient 攻击入口和跨任务结果。",
            ],
        ],
        widths_cm=[3.0, 5.2, 5.2, 4.6],
        font_size=7.8,
    )

    add_heading(doc, "五、问题对应的场景与意义", 1)
    add_body(
        doc,
        "该问题对应的真实场景是多方拥有敏感文本、但又需要共同微调或适配语言模型的协作训练。典型例子包括医院之间联合训练临床文本分类模型、金融机构联合优化舆情或风控文本模型、企业之间共享客服/工单语义能力，以及移动端或边缘端在本地数据上更新个性化语言模型。此时服务端、聚合方或半诚实参与方即使拿不到原文，也可能通过梯度或 adapter 更新进行文本恢复。",
    )
    add_table(
        doc,
        ["应用场景", "为什么会共享更新", "泄露风险", "本工作意义"],
        [
            [
                "跨机构文本联邦学习",
                "机构间数据不能集中，但希望共享模型能力",
                "聚合方或参与方可从 batch 梯度推断训练文本",
                "为 FL-LLM 提供 privacy audit 与 defense pipeline。",
            ],
            [
                "PEFT / LoRA 协作微调",
                "只上传 adapter 或低秩参数更新以降低训练成本",
                "轻量更新仍可能携带样本级 token 线索",
                "验证 defense 是否能贴近未来更常用的轻量微调流程。",
            ],
            [
                "边缘端或个性化模型更新",
                "设备端保留用户输入，周期性上传更新",
                "短文本、小 batch 和高重复 token 会放大恢复风险",
                "提升用户侧文本数据保护能力。",
            ],
            [
                "科研与合规评估",
                "需要定量说明某种训练协议是否安全",
                "仅报告 accuracy 无法说明更新是否泄露文本",
                "建立 privacy-utility-Pareto 和机制消融证据，支撑论文与合规报告。",
            ],
        ],
        widths_cm=[3.5, 4.8, 4.8, 4.8],
        font_size=7.8,
    )
    add_note_box(
        doc,
        "本工作的实际意义",
        [
            "从工程上，它给 FedLLM / PEFT 训练流程提供可复现实验链路：攻击复现、defense sweep、utility 训练、消融和后续 PEFT/partial-gradient 扩展。",
            "从论文上，它把问题从“某个 defense 是否让 DAGER 失败”推进到“中间更新的 recoverability 应该如何被结构性压缩，并在相近 privacy 下保留更多任务信号”。",
            "从当前结果看，Projection-LRB / LRB-lite 是最值得继续推进的主方法候选：它直接对应场景中的低损耗隐私更新需求，也与消融证据一致。",
        ],
        fill="F3F8FB",
    )

    add_heading(doc, "六、威胁模型与评价对象", 1)
    add_body(
        doc,
        "本工作关注大模型联邦训练或分布式微调过程中由中间更新信息导致的数据泄露。客户端不上传原始文本，但会上传梯度、模型更新或 PEFT adapter 更新；攻击者可能据此恢复训练样本。这个问题的关键不在于“数据有没有直接共享”，而在于“共享更新本身是否携带足以恢复原始文本的高分辨率信息”。",
    )
    add_table(
        doc,
        ["攻击面", "代表方法 / 文献", "当前工程状态", "论文作用"],
        [
            ["Full-gradient inversion", "DAGER / LAMP", "DAGER 已成为主实验；n_inputs=100 defense sweep 已完成", "证明 clean FedSGD 梯度存在严重文本恢复风险"],
            ["PEFT leakage", "Gradient Inversion Attacks on PEFT / ReCIT", "LoRA eval-first 框架已有，缺完整结果表", "验证 LRB 是否对实际轻量微调更新也有价值"],
            ["Partial-gradient / layer-level leakage", "Partial Transformer Gradients", "威胁已建模，显式攻击入口仍待补", "证明防御不是只对 full-gradient 攻击有效"],
        ],
        widths_cm=[3.6, 4.6, 5.2, 4.6],
        font_size=8.6,
    )
    add_table(
        doc,
        ["威胁模型要素", "当前实验默认假设", "审稿风险 / 后续补充"],
        [
            [
                "攻击者能力",
                "半诚实服务端或聚合方可观察客户端上传的单轮 full-gradient 更新，并运行 DAGER 类白盒恢复攻击。",
                "需要明确不是 secure aggregation 后的只见总和场景；若考虑多客户端聚合，应另设实验。",
            ],
            [
                "攻击者知识",
                "知道模型结构、当前权重、tokenizer、任务形式和 batch size；可使用同一 fine-tuned checkpoint 进行恢复。",
                "若实际场景中 label、batch size 或客户端采样未知，需要补鲁棒性或敏感性分析。",
            ],
            [
                "被保护对象",
                "客户端本地 batch 中的原始文本 token / sequence，而不是仅保护标签或成员关系。",
                "需要在指标定义里说明 rec_token_mean 与 ROUGE 衡量的是文本内容恢复，不是 DP 隐私损失。",
            ],
            [
                "当前主要攻击面",
                "SST2 + GPT2 fine-tuned checkpoint + batch=2 + n_inputs=100 的 full-gradient DAGER。",
                "PEFT/LoRA 更新、partial-gradient 和跨 backbone 目前是待验证外推，不应写成已有结论。",
            ],
            [
                "防御目标",
                "降低中间更新的 recoverability，同时尽量保持 downstream accuracy / macro-F1。",
                "需要与 topk/compression 在相同训练配置、相同 attack budget 和相同 privacy target 下比较。",
            ],
        ],
        widths_cm=[3.4, 6.8, 6.8],
        font_size=7.6,
    )

    add_heading(doc, "七、已有 baseline 与当前创新位置", 1)
    add_body(
        doc,
        "参考 docs/FL-LLM.md 和 baseline 对比材料，当前 baseline 可以按防御逻辑分成扰动系、隐私标准系、通信压缩系、表示扰动系和训练正则系。它们都能提供必要对照，但大多不是直接以“最小化中间更新的 recoverability”为目标。LRB 的创新位置正在这里：不是粗暴减少数值精度，也不是平均加噪，而是按层和方向建立恢复瓶颈。",
    )
    add_table(
        doc,
        ["baseline", "主要优势", "主要不足", "更适合扮演的角色"],
        [
            ["noise", "实现简单，最小干预", "需要较大噪声才明显有效，utility 快速下降", "sanity-check baseline"],
            ["DP-SGD", "有明确 DP 语义", "当前大模型/FedSGD 设置下 utility 代价很大", "理论标准 baseline"],
            ["topk", "计算/通信成本低，当前 DAGER 下很强", "不是隐私中心设计，泛化到其他攻击面仍需验证", "强通信压缩 baseline"],
            ["compression", "兼顾通信效率和信息损失，工程接入容易", "更多回答少传多少，不直接建模少泄露多少", "强量化 baseline"],
            ["soteria", "直接触及表示泄露", "迁移到 GPT2/LLM 当前设置后稳定性差，privacy 反而恶化", "表示层 baseline / 附录对照"],
            ["mixup", "utility 友好", "当前 DAGER 下 privacy 明显失败", "训练侧弱防御 baseline"],
            ["Projection-LRB", "显式针对 layer-wise recoverability bottleneck；消融显示 projection 是主效应", "多 seed、keep-ratio sweep、PEFT/partial-gradient 证据未补齐", "当前应收束的主方法候选"],
            ["full_lrb", "包含 projection、clipping、residual noise，privacy 很强", "完整配置当前过防御，runtime 偏高", "强防御 / 过防御消融对照"],
        ],
        widths_cm=[2.8, 4.8, 5.4, 4.0],
        font_size=7.9,
    )

    add_heading(doc, "八、Projection-LRB 方法机制", 1)
    add_body(
        doc,
        "根据当前消融证据，文档中的主方法应从泛称 LRB 收束为 Projection-LRB：一个面向 recoverability 的 layer-wise gradient projection defense。full_lrb 仍作为强防御 / 过防御对照保留，用来说明 clipping 与 residual-space noise 并不是当前 full-gradient DAGER 设置下的必要主效应。该实现已经在 train_method=full 的训练 utility 路径中接通，但 LoRA 训练期 defense、forward-side representation bottleneck 和完整 calibration pipeline 仍未落地。",
    )
    add_body(
        doc,
        "更适合论文的方法表述是：对每层梯度 g_l，先根据结构先验与当前 batch 梯度统计估计敏感度 s_l，再用 keep ratio k_l 控制低分辨率 signed_pool 投影 P_{k_l}(g_l)。Projection-LRB 的核心更新可写作 \\tilde{g}_l=P_{k_l}(g_l)，其中高敏感层使用更低分辨率 bottleneck；full_lrb 则在此基础上额外加入 clipping 与 residual-space noise。",
    )
    add_table(
        doc,
        ["模块", "当前实现", "核心作用", "消融关注点"],
        [
            ["Layer-wise sensitivity", "结构先验 + 当前 batch 梯度统计；默认 empirical_weight=0.6", "识别哪些层更可能泄露 token 细节", "rule_only / empirical_only / uniform_all_sensitive"],
            ["Low-resolution signed_pool projection", "先随机符号翻转，再 pooling/interpolation 低分辨率重建", "Projection-LRB 主体；保留粗结构，移除高分辨率样本细节", "proj_only / proj_clip / pool_full"],
            ["Layer-wise clipping", "按所有层范数中位数设置每层裁剪阈值", "full_lrb 中的附加约束；限制个别层的大幅度信息流", "clip_only / proj_clip"],
            ["Residual-space noise", "噪声主要加在低分辨率表示之外的残差方向", "full_lrb 中的强防御组件；进一步污染攻击依赖的细节方向", "full_lrb vs proj_only"],
        ],
        widths_cm=[4.0, 5.3, 4.5, 3.2],
        font_size=8.2,
    )
    add_note_box(
        doc,
        "推荐当前方法表述",
        [
            "Projection-LRB is a layer-wise adaptive gradient projection defense that suppresses recoverability by mapping sensitive-layer gradients into a low-resolution signed pooling subspace.",
            "full_lrb 可作为包含 clipping 与 residual-space noise 的强防御对照；当前不建议把 full_lrb 写成最终主方法。",
        ],
        fill="FFF6E6",
    )

    add_heading(doc, "九、实验设置与数据来源", 1)
    add_table(
        doc,
        ["实验块", "日志目录 / 文档", "关键设置", "用途"],
        [
            [
                "DAGER defense baselines",
                "log/runs/defense_baselines_sst2_b2_gpt2_20260501_010024",
                "sst2, gpt2, ./models/gpt2-ft-rt, batch=2, n_inputs=100",
                "量化 clean 泄露与各 defense 的 attack-time privacy",
            ],
            [
                "Utility focused runs",
                "log/runs/utility260426",
                "full training, epoch=1, seeds=101/202/303",
                "评估 defense 对 accuracy/F1/loss/runtime 的影响",
            ],
            [
                "LRB ablation",
                "log/runs/lrb_ablation_sst2_b2_gpt2_k0.5_20260501_004737",
                "variants=none/identity/clip/proj/full/pool/rule/empirical/uniform, n_inputs=100",
                "拆解 LRB 主效应来自哪个模块",
            ],
            [
                "分析文档",
                "docs/*.md; docs/参考/*",
                "当前工作状态、方法详解、utility、defense、消融、参考图表要求",
                "统一论文叙事、图表清单和下一步路线",
            ],
        ],
        widths_cm=[3.3, 5.0, 5.6, 4.0],
        font_size=8.1,
    )
    add_table(
        doc,
        ["比较口径", "当前做法", "需要在论文中明确的点"],
        [
            [
                "模型与数据",
                "主结果使用 SST2、GPT2 fine-tuned checkpoint、batch=2；clean checkpoint 与 defense runs 对齐。",
                "这是当前强证据边界；跨数据集/backbone 是后续泛化实验。",
            ],
            [
                "攻击预算",
                "DAGER defense baseline sweep 使用 n_inputs=100；同一攻击脚本和同一恢复指标汇总。",
                "不同 defense 的 DAGER=0 只表示在该攻击预算下未恢复，不是不可攻击证明。",
            ],
            [
                "训练效用",
                "utility runs 使用 full training，报告 accuracy、macro-F1、loss、training time。",
                "需要补多 seed 均值/标准差；当前个别负 utility_drop 应按随机波动或 clean-level utility 表述。",
            ],
            [
                "baseline 调参",
                "topk、compression、LRB/full_lrb 均有 sweep 或补点；当前最强 fixed points 是 topk@0.1 与 compression@8。",
                "主表应按同等 privacy target 选各方法最佳 utility 点，避免固定点不公平。",
            ],
            [
                "DP 口径",
                "dpsgd 采用逐样本裁剪 + 高斯噪声的 DP-SGD-style 实现。",
                "未做 privacy accountant 时不能声称 epsilon/delta 形式化 DP 保证。",
            ],
        ],
        widths_cm=[3.2, 6.8, 6.8],
        font_size=7.6,
    )
    add_table(
        doc,
        ["指标", "含义", "解读注意"],
        [
            [
                "rec_token_mean",
                "DAGER 恢复文本与真实文本在 token 级别的平均恢复程度；越低表示当前攻击恢复出的 token 越少。",
                "rec_token_mean=0 表示当前攻击和阈值下未恢复 token，不等价于形式化隐私安全。",
            ],
            [
                "ROUGE-1 + ROUGE-2",
                "恢复文本与真实文本的 n-gram overlap 汇总；用于补充 token-level recovery 的文本相似度视角。",
                "ROUGE 归零与 token 恢复归零共同说明当前攻击失败，但仍需 adaptive attack / 其他攻击面验证。",
            ],
            [
                "eval_accuracy / macro-F1",
                "defense 后完整训练或评估在 SST2 分类任务上的 utility。",
                "小幅高于 clean 的结果应谨慎写作 clean-level utility，需多 seed 支撑显著提升。",
            ],
            [
                "utility_drop",
                "相对 clean accuracy 的下降幅度；越低越好。",
                "负值通常反映随机波动或训练方差，不应直接作为方法提升 accuracy 的证据。",
            ],
            [
                "train_time / attack_time",
                "defense 训练和攻击评估耗时。",
                "用于复杂度与实用性分析，尤其要解释 Projection-LRB 与 full_lrb 的额外开销。",
            ],
        ],
        widths_cm=[3.2, 6.7, 6.9],
        font_size=7.6,
    )

    add_heading(doc, "十、DAGER Defense Baselines：privacy 结果", 1)
    add_body(
        doc,
        "正式 n_inputs=100 结果强化了一个判断：在 SST2+GPT2+batch=2 的 full-gradient DAGER 设置下，clean FedSGD 泄露非常严重，rec_token_mean=0.833506，ROUGE-1+ROUGE-2=141.710856。真正进入 strong privacy baseline 的主要是 LRB、Top-k 和 Compression 三类。",
    )
    add_table(
        doc,
        ["defense", "参数区间 / 关键点", "rec_token_mean", "R1+R2", "结论"],
        summarize_defense_by_family(defense_rows),
        widths_cm=[2.5, 3.4, 3.0, 2.6, 6.8],
        font_size=8.1,
    )
    add_image(doc, charts["privacy"], 16.0, "图 1：代表性 defense 的 DAGER rec_token_mean。")
    add_note_box(
        doc,
        "privacy 口径注意",
        [
            "compression@2 虽然前 56 个样本指标为 0，但本次 n_inputs=100 在第 56/100 处 SVD 失败，正式稳定成功点只应计入 compression@4/8/16。",
            "mixup 与 soteria 在当前设置下不是有效 privacy defense：mixup 多点 rec_token_mean 高于 clean，soteria 几乎完全泄露。",
        ],
    )

    add_heading(doc, "十一、End-to-End Utility：训练效用结果", 1)
    add_body(
        doc,
        "Utility 结果说明，能够把 DAGER 打到 0 并不自动意味着 defense 好。当前 fixed-point 下，topk@0.1 与 compression@8 在保持 DAGER=0 的同时几乎不损失 accuracy；LRB@0.2 也能 DAGER=0，但明显过防御。后来补齐的 full_lrb@0.5 明显改善了 utility，却仍弱于 topk/compression 的强经验点。需要注意的是，单次运行中略高于 clean 的 accuracy 应按 clean-level utility 或训练方差处理，不能直接声称显著提升。",
    )
    util_rows = []
    for row in utility_rows:
        util_rows.append(
            [
                str(row["method"]),
                fmt_float(row["eval_accuracy"], 6),
                fmt_float(row["eval_macro_f1"], 6),
                fmt_float(row["eval_loss"], 6),
                fmt_float(row["utility_drop"], 6),
                str(row["time"]),
                str(row["source"]),
            ]
        )
    add_table(
        doc,
        ["method", "accuracy", "macro-F1", "loss", "utility_drop", "train_time", "source"],
        util_rows,
        widths_cm=[2.8, 2.2, 2.2, 2.1, 2.4, 2.3, 4.3],
        font_size=7.6,
    )
    add_image(doc, charts["utility"], 16.0, "图 2：同等或近似 DAGER privacy 下的 utility drop。")

    add_heading(doc, "十二、Privacy-Utility Tradeoff 与当前 Pareto 判断", 1)
    add_body(
        doc,
        "把 privacy 和 utility 对齐后，当前 full-gradient DAGER 的经验结论很清楚：topk@0.1 与 compression@8 是当前最强 fixed-point baseline；topk@0.3 与 compression@16 已补齐 utility，但没有超过它们；full_lrb@0.5 是当前完整 LRB 配置中最好的 utility 点，但仍有约 2.1 个 accuracy 点的下降。Projection-LRB 的 ablation 点显示了新的主方法空间，但还需要在 main result 中按同等 privacy target 与 topk/compression 同台、多 seed 比较。",
    )
    add_image(doc, charts["scatter"], 16.2, "图 3：privacy-utility tradeoff。左下越好，projection-only 消融点显示新的主方法候选空间。")
    add_table(
        doc,
        ["层级", "方法", "事实依据", "当前判断"],
        [
            ["第一层", "topk@0.1; compression@8", "DAGER=0；accuracy 分别为 0.912462 / 0.911315", "full-gradient DAGER 下当前最强经验 tradeoff"],
            ["第二层", "topk@0.3; compression@16", "DAGER=0；accuracy=0.910933 / 0.909021", "已补齐但未反超强点"],
            ["候选主方法", "proj_only@0.5 / Projection-LRB", "Ablation: DAGER=0；accuracy=0.915520；drop=-0.002294", "单次消融显示 clean-level utility，需纳入正式 main result 和多 seed"],
            ["第三层", "full_lrb@0.5", "DAGER=0；accuracy=0.892584；drop=0.020642", "完整 LRB 当前最好 utility 点，但仍偏重"],
            ["第四层", "lrb@0.2/0.35", "DAGER=0；drop=0.091361 / 0.045107", "privacy 饱和，utility 不足"],
            ["失败/附录层", "mixup/noise/dpsgd/soteria", "要么 privacy 不成立，要么 utility 代价过大", "作为 baseline coverage 保留，不宜做主竞争点"],
        ],
        widths_cm=[2.2, 4.0, 6.0, 5.2],
        font_size=8.2,
    )

    add_heading(doc, "十三、LRB 消融：机制证据", 1)
    add_body(
        doc,
        "完整消融结果是当前最有价值的新证据。它显示 full_lrb 并不是当前最优形式：真正的主效应来自低分辨率 signed_pool projection。proj_only@0.5 在不加 residual noise、不做额外 clipping 的情况下已经把 DAGER rec_token_mean 和 ROUGE 恢复全部降到 0，并在该次运行中保持 clean-level accuracy。这里应谨慎表述为机制线索和主方法候选，而不是已经证明 Projection-LRB 在统计意义上优于所有 baseline。",
    )
    order = ["none", "identity_lrb", "clip_only", "proj_only", "proj_clip", "full_lrb", "pool_full", "rule_only", "empirical_only", "uniform_all_sensitive"]
    ab_map = {r["variant"]: r for r in ablation_rows}
    ab_rows = []
    for v in order:
        if v not in ab_map:
            continue
        r = ab_map[v]
        conclusion = {
            "none": "clean 锚点，泄露严重",
            "identity_lrb": "证明 LRB 管线本身不改变结果",
            "clip_only": "clipping 几乎不能阻断恢复",
            "proj_only": "主效应线索：DAGER=0 且 clean-level utility",
            "proj_clip": "DAGER=0，utility 接近 clean",
            "full_lrb": "DAGER=0，但 residual noise/完整配置过重",
            "pool_full": "普通 pool 弱于 signed_pool",
            "rule_only": "只用规则/完整强防御 utility 代价大",
            "empirical_only": "能防住但有 utility 代价和开销",
            "uniform_all_sensitive": "一刀切 layer-wise 设计不佳",
        }.get(v, "")
        ab_rows.append(
            [
                v,
                fmt_float(r["rec_token_mean"], 6),
                fmt_float(r["r1_plus_r2"], 3),
                fmt_float(r["eval_accuracy"], 6),
                fmt_float(r["utility_drop"], 6),
                seconds_to_hms(r["train_time_seconds"]),
                seconds_to_hms(r["attack_time_seconds"]),
                conclusion,
            ]
        )
    add_table(
        doc,
        ["variant", "rec_token", "R1+R2", "accuracy", "drop", "train_time", "attack_time", "机制结论"],
        ab_rows,
        widths_cm=[3.2, 2.1, 2.0, 2.1, 1.8, 2.2, 2.2, 4.7],
        font_size=7.3,
    )
    add_image(doc, charts["ablation"], 16.0, "图 4：LRB 消融的 utility 对比。")
    add_note_box(
        doc,
        "消融给出的关键改口",
        [
            "主方法身份建议明确为 Projection-LRB / LRB-lite：layer-wise low-resolution signed_pool projection；full_lrb@0.5 不再写成最终主方法。",
            "full_lrb 可保留为强防御 / 过防御对照，用来说明 residual noise 和完整配置并非当前 full-gradient DAGER 的必要主效应。",
            "proj_only@0.5 的负 utility_drop 不应写成显著提升，只能写成该次运行达到 clean-level utility；正式论文主表需要多 seed 均值与标准差。",
        ],
        fill="EAF5F2",
    )

    add_heading(doc, "十四、工程进度与已接通能力", 1)
    add_table(
        doc,
        ["模块", "当前状态", "已完成能力", "主要缺口"],
        [
            ["DAGER baseline", "已接通并有正式结果", "defense_baselines.sh；collect_experiment_logs.py；n_inputs=100 sweep", "跨数据集/backbone 仍待补"],
            ["Full training utility", "已接通", "train.py 支持 full training --defense lrb/topk/compression/noise/dpsgd/mixup", "LRB runtime 和 seed 方差需要优化"],
            ["LRB preset / ablation", "已完成一轮完整消融", "identity/clip/proj/full/pool/rule/empirical/uniform", "proj_only keep-ratio sweep 与 projection-only 细消融待补"],
            ["PEFT / LoRA", "eval-first 最小框架已具备", "加载本地 .pt/.pth LoRA checkpoint；支持 none/noise/topk/compression/lrb", "训练期 LoRA defense、BERT PEFT、Adapter/IA3/Prefix 未接通"],
            ["Partial-gradient", "文档建模完成，代码入口不足", "已有威胁模型和实验计划", "需要 gradient_layer_subset / gradient_param_filter 等入口"],
        ],
        widths_cm=[3.0, 3.2, 6.1, 5.2],
        font_size=8.0,
    )

    add_heading(doc, "十五、论文逻辑与图表规划", 1)
    add_body(
        doc,
        "参考用户提供的三张图，当前材料可以组织成五阶段逻辑：任务价值、机制缺陷、替代机制、技术落地、最终交付。现阶段最需要的是把实验事实转换成一组清晰图表，并避免过度宣称。",
    )
    add_table(
        doc,
        ["阶段", "核心问题", "当前实质内容", "还缺什么证据"],
        [
            ["1", "这件事凭什么值得做？", "FedSGD/LLM clean 梯度可被 DAGER 严重恢复；rec_token=0.833506", "Introduction 中补任务价值与强攻击文献 anchor"],
            ["2", "能力卡在哪个环节？", "泄露来自梯度/更新中的高分辨率 recoverability 结构", "把 gradient 子空间暴露 token 信息讲成明确 problem statement"],
            ["3", "别人为什么绕不开？", "noise/DP utility 代价大；topk/compression 强但不是隐私中心设计", "解释压缩类方法在 full-gradient DAGER 上强但泛化性未证"],
            ["4", "新机制如何落地？", "Projection-LRB：layer-wise sensitivity -> signed_pool bottleneck -> defended update", "keep-ratio sweep 与 PEFT/partial-gradient 验证"],
            ["5", "最终交付什么？", "当前可交付方法候选、消融结论、主结果表雏形", "跨攻击面主表、复杂度分析、稳定结论"],
        ],
        widths_cm=[1.3, 4.0, 6.6, 5.8],
        font_size=8.0,
    )
    add_table(
        doc,
        ["图表类型", "论文位置", "核心功能", "当前状态"],
        [
            ["Main Results 对比表", "Experiments - Main Results", "比较 none/topk/compression/Projection-LRB/full_lrb 等", "已有数据，需纳入 projection 点、多 seed 和强 baseline"],
            ["Ablation Study 表", "Experiments - Ablation", "证明 projection 是主效应", "已具备完整结果"],
            ["Privacy-Utility Pareto 图", "Experiments - Analysis", "同等 privacy 下比较效用", "已有初版，需补 proj_only sweep"],
            ["Framework Overview 图", "Method", "30 秒看懂 Projection-LRB 数据流", "需绘制正式论文图"],
            ["复杂度分析图", "Experiments - Analysis", "展示 accuracy vs train/attack/proxy runtime", "已有时间数据，需整理"],
            ["Transfer / PEFT 表", "Experiments - Transfer", "证明跨攻击面泛化", "尚待实验"],
        ],
        widths_cm=[3.7, 3.8, 6.0, 4.0],
        font_size=8.1,
    )

    add_heading(doc, "十六、当前不应过度声称的内容", 1)
    add_bullets(
        doc,
        [
            "不能写成 LRB/Projection-LRB 已在 SST2/GPT2 上全面优于 topk/compression。当前 full-gradient DAGER 下，topk@0.1 与 compression@8 的经验 tradeoff 更强。",
            "不能写成 full_lrb 是最终主方法。消融已经显示 proj_only@0.5 更像当前主候选，应收束为 Projection-LRB。",
            "不能说 Projection-LRB 在 LoRA/PEFT 或 partial-gradient 下必然最优。这些攻击面还没有完整结果。",
            "不能把 full_lrb 或 Projection-LRB 当成严格 DP 方法。它们有裁剪/投影/加噪组件，但没有完整 epsilon/delta 证明。",
            "不能把 proj_only@0.5 的负 utility_drop 写成显著提升；在多 seed 前只能写成 clean-level utility。",
            "不能把 compression@2 算入稳定成功点。本次 n_inputs=100 在 56/100 处失败。",
        ]
    )

    add_heading(doc, "十七、下一步优先级", 1)
    add_table(
        doc,
        ["优先级", "任务", "目标", "建议输出"],
        [
            ["P0", "把主方法身份统一为 Projection-LRB / proj_only@0.5", "让论文方法与消融证据一致", "主方法描述、算法伪代码、main table 新行"],
            ["P0", "proj_only keep-ratio sweep + 多 seed", "找 DAGER=0 时 utility 最高的最宽松 bottleneck，并估计方差", "k=0.5/0.65/0.75/0.9 曲线；mean±std 表"],
            ["P0", "projection-only 细消融", "拆开 rule / empirical / uniform / no_empirical 的作用", "proj_rule_only、proj_empirical_only、proj_uniform、proj_no_empirical 表"],
            ["P1", "重做主结果和 Pareto 表", "把 topk/compression/Projection-LRB/full_lrb 同台比较", "main results + Pareto 图"],
            ["P1", "LoRA/PEFT 对照", "验证跨实际轻量更新攻击面", "none/proj_only/proj_clip/full_lrb/topk/compression 的 PEFT 表"],
            ["P2", "partial-gradient / layer-level leakage", "验证局部更新泄露下的结构性防御价值", "first block / qkv / last layers 等攻击面表"],
            ["P2", "跨数据集/backbone", "避免单一 SST2/GPT2 偶然性", "cola/rte/gpt2; sst2/bert 等结果"],
            ["P2", "runtime 优化", "降低 Projection-LRB 和 full_lrb 实用开销", "train time / attack time / proxy runtime 分析图"],
        ],
        widths_cm=[1.7, 4.6, 6.6, 4.8],
        font_size=8.0,
    )

    add_heading(doc, "十八、可以直接用于汇报的阶段性结论", 1)
    add_note_box(
        doc,
        "中文结论",
        [
            "当前 SST2+GPT2+batch=2 的 full-gradient DAGER 实验表明，clean FedSGD 梯度存在严重文本恢复风险。full_lrb 完整配置能稳定把 DAGER 恢复压到 0，但其 clipping 和 residual noise 带来了明显 utility 代价。消融进一步显示，低分辨率 signed_pool projection 是当前 LRB 有效性的主因：proj_only@0.5 在不加噪、不做完整裁剪的情况下即可实现 DAGER=0，并在单次运行中保持 clean-level accuracy。因此，下一阶段应把主方法身份统一为 Projection-LRB / LRB-lite，并通过 keep-ratio sweep、多 seed、LoRA/PEFT、partial-gradient 和跨数据集实验验证其泛化价值。",
        ],
        fill="EAF5F2",
    )
    add_note_box(
        doc,
        "English statement",
        [
            "In the current SST2/GPT2 full-gradient DAGER setting, the dominant effective component is the low-resolution signed projection bottleneck. While the full LRB configuration suppresses token recovery to zero, its clipping and residual-space noise introduce unnecessary utility loss. The projection-only variant achieves zero token and ROUGE recovery with clean-level utility in the current ablation run, making Projection-LRB the most promising method candidate pending multi-seed and cross-attack validation.",
        ],
        fill="F3F8FB",
    )

    add_heading(doc, "十九、参考文献与 baseline 对应关系", 1)
    add_body(
        doc,
        "本节先按当前工作实际用途整理参考文献：第一类用于支撑 DAGER 等文本恢复攻击和实验框架，第二类对应已跑的 defense baselines，第三类支撑 FL-LLM、PEFT/LoRA 与 partial-gradient 等问题场景。后续写论文时可再统一转换为 BibTeX 或 GB/T 7714 格式。",
    )
    add_table(
        doc,
        ["类别", "参考文献", "与本文工作的关系"],
        [
            [
                "主攻击框架",
                "Petrov et al. DAGER: Exact Gradient Inversion for Large Language Models. NeurIPS 2024.",
                "当前 privacy 实验的主攻击与主评测框架；用于证明 full-gradient LLM/FedSGD 设置下文本可被高精度恢复。",
            ],
            [
                "文本梯度泄露",
                "Balunovic et al. LAMP: Extracting Text from Gradients with Language Model Priors. NeurIPS 2022.",
                "DAGER 之前的重要文本恢复攻击；用于铺垫语言模型先验可显著增强梯度反演。",
            ],
            [
                "Transformer 梯度攻击",
                "Deng et al. TAG: Gradient Attack on Transformer-based Language Models. Findings of EMNLP 2021.",
                "说明 Transformer/NLP 模型梯度泄露并非偶然现象，是 DAGER/LLM 攻击线的早期代表。",
            ],
            [
                "通用梯度反演",
                "Zhu, Liu, and Han. Deep Leakage from Gradients. NeurIPS 2019.",
                "通用 gradient inversion 起点；用于介绍 federated learning 中梯度本身可能泄露训练样本。",
            ],
            [
                "通用优化式反演",
                "Geiping et al. Inverting Gradients: How easy is it to break privacy in federated learning? NeurIPS 2020.",
                "补充传统优化式反演背景，说明该问题不局限于文本任务。",
            ],
        ],
        widths_cm=[3.0, 7.2, 7.0],
        font_size=7.6,
    )
    add_table(
        doc,
        ["当前 baseline / 扩展点", "参考文献", "放入文档时的定位"],
        [
            [
                "DP-SGD / dpsgd",
                "Abadi et al. Deep Learning with Differential Privacy. ACM CCS 2016.",
                "对应逐样本裁剪 + 高斯噪声的理论隐私 baseline；当前实验中 utility 代价较高。",
            ],
            [
                "noise",
                "Abadi et al. Deep Learning with Differential Privacy. ACM CCS 2016; Gaussian noise baseline as a simple perturbation control.",
                "作为最朴素的梯度扰动对照，不单独声称具备 DP 保证。",
            ],
            [
                "topk",
                "Aji and Heafield. Sparse Communication for Distributed Gradient Descent. EMNLP 2017; Lin et al. Deep Gradient Compression. ICLR 2018.",
                "定位为通信压缩/稀疏化强经验 baseline；当前 full-gradient DAGER 下 tradeoff 很强，但不是隐私中心设计。",
            ],
            [
                "compression",
                "Alistarh et al. QSGD: Communication-Efficient SGD via Gradient Quantization and Encoding. NeurIPS 2017.",
                "对应当前代码中的 QSGD-style stochastic quantization；用于量化压缩类 baseline。",
            ],
            [
                "soteria",
                "Sun et al. Soteria: Provable Defense Against Privacy Leakage in Federated Learning from Representation Perspective. CVPR 2021.",
                "表示层隐私防御 baseline；当前 GPT2/DAGER 设置下效果不好，应作为对照而非主竞争点。",
            ],
            [
                "mixup",
                "Zhang et al. mixup: Beyond Empirical Risk Minimization. ICLR 2018.",
                "训练侧增强/混合样本 baseline；当前结果 utility 友好但 privacy 失败。",
            ],
        ],
        widths_cm=[3.4, 7.8, 6.2],
        font_size=7.4,
    )
    add_table(
        doc,
        ["用途", "参考文献", "建议写法"],
        [
            [
                "联邦学习场景",
                "McMahan et al. Communication-Efficient Learning of Deep Networks from Decentralized Data. AISTATS 2017.",
                "用于介绍 FedAvg/FedSGD 和去中心化数据训练背景。",
            ],
            [
                "LoRA / PEFT 场景",
                "Hu et al. LoRA: Low-Rank Adaptation of Large Language Models. ICLR 2022.",
                "用于说明只共享 LoRA/adapter 更新的现实训练流程。",
            ],
            [
                "PEFT 梯度反演",
                "Sami et al. Gradient Inversion Attacks on Parameter-Efficient Fine-Tuning. CVPR 2025.",
                "用于支撑 PEFT 更新同样可能泄露数据；建议放在相关工作或未来扩展。",
            ],
            [
                "LLM PEFT 私有数据恢复",
                "Xie et al. ReCIT: Reconstructing Full Private Data from Gradient in Parameter-Efficient Fine-Tuning of Large Language Models. arXiv 2025.",
                "预印本，适合作为 PEFT leakage 风险补充，不建议写成已发表顶会结论。",
            ],
            [
                "Partial-gradient 泄露",
                "Li et al. Seeing the Forest through the Trees: Data Leakage from Partial Transformer Gradients. EMNLP 2024.",
                "支撑下一步 partial-gradient 泛化实验，说明只暴露部分 Transformer 梯度也可能泄露数据。",
            ],
        ],
        widths_cm=[3.0, 7.8, 6.4],
        font_size=7.4,
    )
    add_note_box(
        doc,
        "引用口径建议",
        [
            "DAGER、LAMP、TAG、DLG 是攻击背景的主线；DAGER 应作为当前实验框架的核心引用。",
            "topk/compression 应明确写作通信压缩类强 baseline，而不是传统隐私防御；它们在当前 DAGER 下强，正好构成本文必须正面比较的经验对手。",
            "ReCIT 目前按 arXiv 预印本处理；PEFT/partial-gradient 文献更适合放在动机、相关工作和未来验证部分。",
        ],
        fill="FFF6E6",
    )

    add_heading(doc, "二十、材料来源清单", 1)
    add_table(
        doc,
        ["类别", "路径"],
        [
            ["模板", str(DOCS / "初步思路模版.docx")],
            ["当前工作分析", str(DOCS / "CURRENT_WORK_STATUS_ANALYSIS_20260427.md")],
            ["DAGER baseline 分析", str(DOCS / "DEFENSE_BASELINES_N100_ANALYSIS_20260502.md")],
            ["Utility 分析", str(DOCS / "UTILITY_RESULTS_ANALYSIS_20260426.md")],
            ["LRB 消融分析", str(DOCS / "LRB_ABLATION_ANALYSIS_20260503.md")],
            ["LRB 方法详解", str(DOCS / "LRB_方法详解.md")],
            ["FL-LLM 原始/整理思路", str(DOCS / "FL-LLM.md")],
            ["PEFT 框架", str(DOCS / "PEFT_EVAL.md")],
            ["参考材料", str(DOCS / "参考")],
            ["DAGER runs", str(RUNS / "defense_baselines_sst2_b2_gpt2_20260501_010024")],
            ["Utility runs", str(RUNS / "utility260426")],
            ["Ablation runs", str(RUNS / "lrb_ablation_sst2_b2_gpt2_k0.5_20260501_004737")],
        ],
        widths_cm=[4.0, 13.0],
        font_size=8.0,
    )

    # Light footer metadata
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("FedLLM / LRB 阶段性工作进展汇总 - 2026-05-05")
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string("6b7280")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
