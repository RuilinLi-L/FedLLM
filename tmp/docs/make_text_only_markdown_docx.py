from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:/code/Projects/FedLLM")
OUT_DIR = ROOT / "output" / "doc"
SRC_DOCX = OUT_DIR / "FedLLM_LRB_阶段性工作进展与实验结果汇总_审稿完善版_20260505.docx"
OUT_MD = OUT_DIR / "FedLLM_LRB_阶段性工作进展与实验结果汇总_纯文字表格版_20260505.md"
OUT_DOCX = OUT_DIR / "FedLLM_LRB_阶段性工作进展与实验结果汇总_纯文字表格版_20260505.docx"


def iter_block_items(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield "p", child
        elif child.tag.endswith("}tbl"):
            yield "tbl", child


def paragraph_from_element(doc: Document, element):
    for p in doc.paragraphs:
        if p._element is element:
            return p
    return None


def table_from_element(doc: Document, element):
    for t in doc.tables:
        if t._element is element:
            return t
    return None


def clean_text(text: str) -> str:
    return " ".join(str(text).replace("\n", " ").replace("\r", " ").split())


def escape_md_cell(text: str) -> str:
    text = clean_text(text)
    text = text.replace("\\", "\\\\").replace("|", "\\|")
    return text if text else " "


def md_table_from_docx_table(table) -> list[str]:
    if not table.rows:
        return []
    rows = [[clean_text(cell.text) for cell in row.cells] for row in table.rows]
    max_cols = max(len(row) for row in rows)
    rows = [row + [""] * (max_cols - len(row)) for row in rows]
    header = rows[0]
    out = [
        "| " + " | ".join(escape_md_cell(cell) for cell in header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in rows[1:]:
        out.append("| " + " | ".join(escape_md_cell(cell) for cell in row) + " |")
    return out


def paragraph_to_md(p) -> list[str]:
    text = clean_text(p.text)
    if not text:
        return []
    style = p.style.name if p.style is not None else ""
    if style == "Title":
        return [f"# {text}"]
    if style.startswith("Heading 1"):
        return [f"## {text}"]
    if style.startswith("Heading 2"):
        return [f"### {text}"]
    if style.startswith("Heading 3"):
        return [f"#### {text}"]
    if style.startswith("List"):
        return [f"- {text}"]
    if text.startswith("图 "):
        return [f"**{text}**"]
    return [text]


def build_markdown() -> str:
    src = Document(str(SRC_DOCX))
    lines: list[str] = []
    for kind, element in iter_block_items(src):
        if kind == "p":
            p = paragraph_from_element(src, element)
            if p is None:
                continue
            md_lines = paragraph_to_md(p)
            if md_lines:
                lines.extend(md_lines)
                lines.append("")
        elif kind == "tbl":
            table = table_from_element(src, element)
            if table is None:
                continue
            md_lines = md_table_from_docx_table(table)
            if md_lines:
                lines.extend(md_lines)
                lines.append("")
    text = "\n".join(lines).strip() + "\n"
    # Remove image-only captions because the pure text version intentionally drops images.
    filtered: list[str] = []
    for line in text.splitlines():
        if line.startswith("**图 ") and line.endswith("**"):
            continue
        filtered.append(line)
    return "\n".join(filtered).strip() + "\n"


def set_doc_defaults(doc: Document):
    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        st = styles[name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.size = Pt(20)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True


def add_code_line(doc: Document, line: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(line)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(8.6)


def build_docx(markdown: str):
    doc = Document()
    set_doc_defaults(doc)
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_paragraph(line[2:], style="Title")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:], style="Heading 3")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("| "):
            add_code_line(doc, line)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(line)
            r.font.name = "微软雅黑"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            r.font.size = Pt(10.5)
    doc.save(OUT_DOCX)


def main():
    markdown = build_markdown()
    OUT_MD.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
